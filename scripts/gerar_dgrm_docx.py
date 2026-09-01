import sys, json, os
from docx import Document
from docx.shared import Pt
from copy import deepcopy

def fill_template(template_path, output_path, data):
    doc = Document(template_path)

    # Replace paragraphs
    replacements = {
        5: f"Fabricante: {data.get('brand', '')}",
        6: f"Manufacturer: {data.get('brand', '')}",
        9: f"Tipo: {data.get('model', '')}  N.º de série: {data.get('serial', '')}  Para: {data.get('capacity', '')} Pessoas",
        10: f"Type: {data.get('model', '')}  Serial nº: {data.get('serial', '')}  For: {data.get('capacity', '')} Persons",
    }

    # Garrafa de gás
    cyl = data.get('cylinderSerial', '')
    def fmt_peso(v):
        try:
            n = float(str(v).replace(',', '.'))
            return f"{n:.3f}"
        except (ValueError, TypeError):
            return str(v or '')
    peso = fmt_peso(data.get('cylinderPesoBruto', ''))
    tara = fmt_peso(data.get('cylinderTara', ''))
    co2 = fmt_peso(data.get('cylinderCo2', ''))
    n2 = fmt_peso(data.get('cylinderN2', ''))
    replacements[12] = f"Garrafa de gás n.º: {cyl}  Peso: {peso}  Tara: {tara}  CO2: {co2}  N2: {n2}"
    replacements[13] = f"Gas Cylinder: {cyl}  Full weight: {peso}  Tare weight: {tara}  CO2: {co2}  N2: {n2}"

    # Pack
    pack = data.get('packType', '')
    replacements[15] = f"Embalagem de Sobrevivência Tipo: {pack}"
    replacements[16] = f"Emergency Pack Type: {pack}"

    # Container
    container = data.get('containerModel', '')
    replacements[18] = f"Contentor/Saco modelo: {container}"
    replacements[19] = f"Container/Valise Type: {container}"

    # Painter line
    painter = data.get('painterLength', '')
    replacements[21] = f"Comprimento do cabo de disparo: {painter}"
    replacements[22] = f"Painter line length: {painter}"

    # HRU
    hru = data.get('hruReferencia', '')
    hruTipo = data.get('cylinderCabecaDisparoRef', '')
    replacements[24] = f"Libertador hidrostático: {hru}  Tipo: {hruTipo}"
    replacements[25] = f"Hydrostatic release: {hru}  Type: {hruTipo}"

    # Data and place
    replacements[28] = f"Local e data de emissão da ficha: _________________________ , _________________________"
    replacements[31] = "O Responsável pela Estação de Serviço:"
    replacements[33] = "________________________________"

    for idx, text in replacements.items():
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            # Clear existing runs and add new text
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)

    # Fill table - Row 0: manufacture date, place, cert number
    t = doc.tables[0]
    fab = data.get('dataFabrico', '')
    cert = data.get('ultimoCertificadoNumero', '') or data.get('certificadoExternoNumero', '') or 'AZ26-169'
    local = data.get('ilha', '') or 'Ponta Delgada'

    # Cell 0,0: Data de fabrico
    cell_00 = t.rows[1].cells[0]
    for p in cell_00.paragraphs:
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = fab
        else: p.add_run(fab)

    # Cell 0,1: Local
    cell_01 = t.rows[1].cells[1]
    for p in cell_01.paragraphs:
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = local
        else: p.add_run(local)

    # Cell 0,2: N.º do certificado original
    cell_02 = t.rows[1].cells[2]
    for p in cell_02.paragraphs:
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = cert
        else: p.add_run(cert)

    # Fill table - Row 1 (Index 3): Data revisão periódica, Navio, N.º do relatório
    data_rev = data.get('dataInspecao', '') or '01/09/2026'
    navio_nome = data.get('shipName', '') or data.get('shipNameManual', '') or 'Rainha da Calheta'
    relatorio_num = data.get('ultimoCertificadoNumero', '') or 'AZ26-169'

    if len(t.rows) > 3:
        row_rev = t.rows[3]
        # Data da revisão periódica
        cell_10 = row_rev.cells[0]
        for p in cell_10.paragraphs:
            for r in p.runs: r.text = ''
            if p.runs: p.runs[0].text = data_rev
            else: p.add_run(data_rev)

        # Navio
        cell_11 = row_rev.cells[1]
        for p in cell_11.paragraphs:
            for r in p.runs: r.text = ''
            if p.runs: p.runs[0].text = navio_nome
            else: p.add_run(navio_nome)

        # N.º do relatório
        cell_12 = row_rev.cells[2]
        for p in cell_12.paragraphs:
            for r in p.runs: r.text = ''
            if p.runs: p.runs[0].text = relatorio_num
            else: p.add_run(relatorio_num)

    doc.save(output_path)
    print(f"DOCX gerado: {output_path}")

if __name__ == '__main__':
    template = sys.argv[1]
    output = sys.argv[2]
    json_path = sys.argv[3]
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    fill_template(template, output, data)
